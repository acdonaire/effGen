"""
DOCX parsing tool for the effGen framework.

Backend: python-docx.

Operations
----------
- text        : Extract all text as a single string.
- paragraphs  : Return structured paragraph list with style info.
- tables      : Extract all tables as list-of-rows.
- metadata    : Return core document properties (author, title, etc.).

Input: file path (str/Path) OR raw DOCX bytes.
Output: structured dicts with success/data/error keys.
"""

from __future__ import annotations

import asyncio
import io
import logging
from pathlib import Path
from typing import Any

from effgen.errors import CorruptDocumentError

from ..base_tool import (
    BaseTool,
    ParameterSpec,
    ParameterType,
    ToolCategory,
    ToolMetadata,
)
from ._fs import confine_path, normalize_allowed_dirs

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Source resolution
# ---------------------------------------------------------------------------

def _resolve_docx(source: str | bytes | Path) -> bytes:
    if isinstance(source, str | Path):
        p = Path(source)
        if not p.exists():
            raise FileNotFoundError(f"DOCX not found: {source}")
        return p.read_bytes()
    if isinstance(source, bytes | bytearray):
        return bytes(source)
    raise TypeError(f"Expected str, Path, or bytes; got {type(source)}")


def _open_docx(raw: bytes):
    """Open a python-docx Document from raw bytes, raising CorruptDocumentError on failure."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCXTool. "
            "Install with: pip install 'effgen[documents]'"
        ) from exc
    try:
        return Document(io.BytesIO(raw))
    except Exception as exc:
        raise CorruptDocumentError("docx", str(exc)) from exc


# ---------------------------------------------------------------------------
# Operation implementations
# ---------------------------------------------------------------------------

def _docx_text(source: str | bytes | Path) -> dict[str, Any]:
    raw = _resolve_docx(source)
    doc = _open_docx(raw)
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    return {
        "success": True,
        "data": {"text": full_text, "paragraph_count": len(paragraphs)},
        "text": full_text,
        "paragraph_count": len(paragraphs),
        "error": None,
    }


def _docx_paragraphs(source: str | bytes | Path) -> dict[str, Any]:
    raw = _resolve_docx(source)
    doc = _open_docx(raw)
    result: list[dict] = []
    for idx, para in enumerate(doc.paragraphs):
        result.append({
            "index": idx,
            "text": para.text,
            "style": para.style.name if para.style else None,
            "bold": any(run.bold for run in para.runs),
            "italic": any(run.italic for run in para.runs),
            "alignment": str(para.alignment) if para.alignment else None,
        })
    return {
        "success": True,
        "data": {"paragraphs": result, "count": len(result)},
        "paragraphs": result,
        "count": len(result),
        "error": None,
    }


def _docx_tables(source: str | bytes | Path) -> dict[str, Any]:
    raw = _resolve_docx(source)
    doc = _open_docx(raw)
    all_tables: list[dict] = []
    for t_idx, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        all_tables.append({
            "table_index": t_idx,
            "rows": rows,
            "row_count": len(rows),
            "col_count": max((len(r) for r in rows), default=0),
        })
    return {
        "success": True,
        "data": {"tables": all_tables, "table_count": len(all_tables)},
        "tables": all_tables,
        "table_count": len(all_tables),
        "error": None,
    }


def _docx_metadata(source: str | bytes | Path) -> dict[str, Any]:
    raw = _resolve_docx(source)
    doc = _open_docx(raw)
    props = doc.core_properties
    info: dict[str, Any] = {
        "author": props.author,
        "title": props.title,
        "subject": props.subject,
        "comments": getattr(props, "comments", None),
        "keywords": props.keywords,
        "last_modified_by": props.last_modified_by,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "revision": props.revision,
        "category": props.category,
        "content_status": props.content_status,
    }
    return {
        "success": True,
        "data": info,
        "metadata": info,
        "error": None,
    }


# ---------------------------------------------------------------------------
# BaseTool subclass
# ---------------------------------------------------------------------------

class DOCXTool(BaseTool):
    """
    Parse Microsoft Word DOCX documents.

    Backend: python-docx.

    Accepts file paths (str/Path) OR raw DOCX bytes.
    Raises CorruptDocumentError on malformed input.
    """

    def __init__(self, allowed_directories: list[str] | None = None) -> None:
        """Args:
            allowed_directories: Roots a file path may be read from. By default
                any path is allowed except protected system and credential
                locations (/etc, /proc, ~/.ssh, cloud creds, …), which are
                always refused. Pass a list to confine reads to those roots only.
        """
        self._allowed_dirs = normalize_allowed_dirs(allowed_directories)
        super().__init__(
            metadata=ToolMetadata(
                name="docx",
                description=(
                    "Parse DOCX (Microsoft Word) documents. Operations: text (full text), "
                    "paragraphs (structured paragraph list with styles), "
                    "tables (extract all tables), metadata (document properties). "
                    "Accepts file paths or raw bytes."
                ),
                category=ToolCategory.DATA_PROCESSING,
                parameters=[
                    ParameterSpec(
                        name="operation",
                        type=ParameterType.STRING,
                        description="Operation: 'text', 'paragraphs', 'tables', or 'metadata'.",
                        required=True,
                        enum=["text", "paragraphs", "tables", "metadata"],
                    ),
                    ParameterSpec(
                        name="path",
                        type=ParameterType.STRING,
                        description="Path to the DOCX file.",
                        required=False,
                    ),
                ],
                timeout_seconds=30,
                tags=["docx", "word", "document", "parsing", "tables"],
                examples=[
                    {"operation": "text", "path": "/path/to/doc.docx"},
                    {"operation": "paragraphs", "path": "/path/to/doc.docx"},
                    {"operation": "tables", "path": "/path/to/doc.docx"},
                    {"operation": "metadata", "path": "/path/to/doc.docx"},
                ],
            )
        )

    async def _execute(
        self,
        operation: str,
        path: str | None = None,
        docx_bytes: bytes | None = None,
        **kwargs: Any,
    ) -> dict[str, Any]:
        source: str | bytes | Path
        if path:
            source = str(confine_path(path, self._allowed_dirs))
        elif docx_bytes is not None:
            source = docx_bytes
        else:
            raise ValueError("Provide 'path' or 'docx_bytes'.")

        op = operation.lower()

        if op == "text":
            return await asyncio.to_thread(_docx_text, source)
        elif op == "paragraphs":
            return await asyncio.to_thread(_docx_paragraphs, source)
        elif op == "tables":
            return await asyncio.to_thread(_docx_tables, source)
        elif op == "metadata":
            return await asyncio.to_thread(_docx_metadata, source)
        else:
            raise ValueError(
                f"Unknown operation: {operation!r}. "
                "Use 'text', 'paragraphs', 'tables', or 'metadata'."
            )
