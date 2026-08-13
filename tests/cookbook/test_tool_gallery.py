"""Snippet tests for ``docs/tools/gallery.md``.

Every Python block in the gallery is checked for form (compiles; ``execute()``
called with keyword arguments only, never a positional dict), and the runnable
ones are executed verbatim in a subprocess:

* offline snippets run everywhere (file paths under ``/tmp/`` are rewritten to
  a per-test temp directory and any input fixture is created there first);
* snippets that reach a public network service are marked ``live``;
* snippets that need credentials, a webhook URL, or a system binary are
  skipped unless that requirement is present.

Each gallery ``###`` heading must be categorized below — an uncategorized
snippet fails ``test_every_gallery_snippet_is_categorized`` so new entries
cannot silently go untested.
"""

from __future__ import annotations

import ast
import os
import shutil
import subprocess
import sys
import time
from pathlib import Path

import pytest
from dotenv import load_dotenv

_REPO_ROOT = Path(__file__).parent.parent.parent
_GALLERY = _REPO_ROOT / "docs/tools/gallery.md"
_AUDIO_FIXTURE = _REPO_ROOT / "tests/fixtures/multimodal/sample_audio.mp3"

# override=False, like every other module that reads this file: without it the
# repository .env wins over a key the operator deliberately exported, so a run
# aimed at one account silently goes to another.
load_dotenv(_REPO_ROOT / ".env", override=False)


# ---------------------------------------------------------------------------
# Gallery parsing
# ---------------------------------------------------------------------------

def _gallery_snippets() -> list[tuple[str, str]]:
    """Return ``(heading, code)`` for every ```python block, in file order.

    The heading is the nearest preceding ``###`` title; headings repeat when a
    tool has several blocks (or appears in two sections), so consumers index
    by position within a heading where it matters.
    """
    text = _GALLERY.read_text(encoding="utf-8")
    snippets: list[tuple[str, str]] = []
    heading = ""
    in_block = False
    lines: list[str] = []
    for line in text.splitlines():
        if line.startswith("### ") and not in_block:
            heading = line[4:].strip()
        elif line.startswith("## ") and not in_block:
            heading = line[3:].strip()
        elif line.strip() == "```python" and not in_block:
            in_block = True
            lines = []
        elif line.strip() == "```" and in_block:
            in_block = False
            snippets.append((heading, "\n".join(lines)))
        elif in_block:
            lines.append(line)
    return snippets


def _blocks_for(heading: str) -> list[str]:
    return [code for h, code in _gallery_snippets() if h == heading]


# ---------------------------------------------------------------------------
# Categorization: every heading with at least one snippet appears exactly once
# ---------------------------------------------------------------------------

# Runnable offline with no credentials; fixture files are created on the fly.
# (QRGenerateTool/QRReadTool are also offline but form a two-step sequence —
# generate writes the PNG that read decodes — so they have a dedicated test.)
OFFLINE = [
    "Calculator",
    "PythonREPL",
    "CodeExecutor",
    "BashTool",
    "FileOps",
    "DataFrameTool",
    "PlotTool",
    "StatsTool",
    "GitTool",
    "SystemInfoTool",
    "JSONTool",
    "DateTimeTool",
    "TextProcessingTool",
    "LanguageDetectTool",
    "EmailDraftTool",
    "SlackDraftTool",
    "ImageInfoTool",
    "PDFTool",
    "DOCXTool",
    "ExcelTool",
]

# Reach a public, unauthenticated network service.
NETWORK = [
    "WebSearch",
    "URLFetch",
    "WikipediaTool",
    "HTTPTool",
    "PubMedTool",
    "ArXivTool",
    "SemanticScholarTool",
    "RSSFeedTool",
    "NewsTool",
    "RedditTool",
    "HackerNewsTool",
    "TranslateTool",
    "StockPriceTool",
    "CurrencyConverterTool",
    "CryptoTool",
    "StackOverflowTool",
    "GitHubTool",
    "WeatherTool",        # appears twice (Utilities + Geo/Weather); both run
    "GeocodeTool",
    "MapsTool",
    "YouTubeTranscriptTool",
    "YouTubeMetadataTool",
]

# Need a credential, webhook URL, system binary, or model API key; each has a
# dedicated test below with its own skip condition.
GATED = [
    "DockerTool",
    "WolframAlphaTool",
    "EmailSMTPTool",
    "EmailIMAPTool",
    "SlackWebhookTool",
    "DiscordWebhookTool",
    "OCRTool",
    "AudioTranscribeTool",
    "ImageCaptionTool",
    "OpenAI Native Tools",
    "Gemini Native Tools",
    "Using Tools in an Agent",
    "Using Presets",
]

# Offline two-step sequences with their own dedicated test.
SEQUENCES = ["QRGenerateTool", "QRReadTool"]

# Optional third-party imports per heading; missing ones skip the test.
_OPTIONAL_DEPS = {
    "QRGenerateTool": ["qrcode"],
    "QRReadTool": ["pyzbar"],
    "DataFrameTool": ["pandas"],
    "PlotTool": ["matplotlib"],
    "LanguageDetectTool": ["langdetect"],
    "ImageInfoTool": ["PIL"],
    "PDFTool": ["pypdf", "matplotlib"],
    "DOCXTool": ["docx"],
    "ExcelTool": ["openpyxl"],
    "MapsTool": ["staticmap"],
    "YouTubeTranscriptTool": ["youtube_transcript_api"],
    "YouTubeMetadataTool": ["yt_dlp"],
    "StockPriceTool": ["yfinance"],
    "RSSFeedTool": ["feedparser"],
}


def test_every_gallery_snippet_is_categorized():
    headings = {h for h, _ in _gallery_snippets()}
    categorized = set(OFFLINE) | set(NETWORK) | set(GATED) | set(SEQUENCES)
    missing = headings - categorized
    assert not missing, f"Uncategorized gallery snippets: {sorted(missing)}"
    stale = categorized - headings
    assert not stale, f"Categorized but absent from the gallery: {sorted(stale)}"


# ---------------------------------------------------------------------------
# Form: every block compiles and calls execute() with keywords only
# ---------------------------------------------------------------------------

def test_gallery_snippets_compile_and_use_keyword_execute():
    problems: list[str] = []
    for heading, code in _gallery_snippets():
        try:
            tree = ast.parse(code)
        except SyntaxError as exc:
            problems.append(f"{heading}: does not parse: {exc}")
            continue
        for node in ast.walk(tree):
            if (
                isinstance(node, ast.Call)
                and isinstance(node.func, ast.Attribute)
                and node.func.attr == "execute"
            ):
                if node.args:
                    problems.append(
                        f"{heading}: execute() called with positional argument(s); "
                        "the tool API takes keyword arguments"
                    )
            if (
                isinstance(node, ast.Subscript)
                and isinstance(node.value, ast.Name)
                and node.value.id == "result"
            ):
                problems.append(
                    f"{heading}: subscripts `result[...]`; ToolResult exposes "
                    ".success/.output/.error attributes"
                )
    assert not problems, "\n".join(problems)


def _result_names(tree: ast.Module) -> set[str]:
    """Names bound to the result of an ``execute()`` call in *tree*."""
    names: set[str] = set()
    for node in ast.walk(tree):
        if not isinstance(node, ast.Assign):
            continue
        calls = [
            n for n in ast.walk(node.value)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "execute"
        ]
        if not calls:
            continue
        for target in node.targets:
            if isinstance(target, ast.Name):
                names.add(target.id)
    return names


def _guard_line(tree: ast.Module, name: str) -> int | None:
    """Line of the earliest ``if not <name>.success:`` that stops the snippet."""
    lines: list[int] = []
    for node in ast.walk(tree):
        if not isinstance(node, ast.If):
            continue
        test = node.test
        if not (
            isinstance(test, ast.UnaryOp)
            and isinstance(test.op, ast.Not)
            and isinstance(test.operand, ast.Attribute)
            and test.operand.attr == "success"
            and isinstance(test.operand.value, ast.Name)
            and test.operand.value.id == name
        ):
            continue
        # The guard has to end the snippet, not just print a note: a body that
        # falls through leaves the following line to subscript ``None``. Only
        # the guard's own body counts — a `raise` in its `else` branch runs
        # when the call succeeded.
        stops = any(
            isinstance(stmt, ast.Raise)
            or (
                isinstance(stmt, ast.Expr)
                and isinstance(stmt.value, ast.Call)
                and isinstance(stmt.value.func, ast.Attribute)
                and stmt.value.func.attr == "exit"
            )
            for body_stmt in node.body
            for stmt in ast.walk(body_stmt)
        )
        if stops:
            lines.append(node.lineno)
    return min(lines) if lines else None


def _first_output_line(tree: ast.Module, name: str) -> int | None:
    """Line of the first read of ``<name>.output``."""
    lines = [
        node.lineno
        for node in ast.walk(tree)
        if isinstance(node, ast.Attribute)
        and node.attr == "output"
        and isinstance(node.value, ast.Name)
        and node.value.id == name
    ]
    return min(lines) if lines else None


# Headings under GATED whose tool also calls a remote host, so their snippets
# carry the same guard as the unauthenticated ones: Wolfram Alpha's API, an
# OCR.space fallback, the vision model router, an SMTP/IMAP server, and the
# Slack and Discord webhook endpoints.
CREDENTIALED_REMOTE = [
    "WolframAlphaTool",
    "OCRTool",
    "ImageCaptionTool",
    "EmailSMTPTool",
    "EmailIMAPTool",
    "SlackWebhookTool",
    "DiscordWebhookTool",
]

# Headings whose example names a file the reader has to supply. The path in the
# snippet does not exist on a reader's machine, so running it as written is a
# failure outside the snippet exactly as an outage is: the tool reports the
# missing path and the read of ``.output`` would discard that report.
READER_SUPPLIED_PATH = [
    "DataFrameTool",
    "AudioTranscribeTool",
    "ImageInfoTool",
    "PDFTool",
    "DOCXTool",
    "ExcelTool",
]

_REMOTE = frozenset(NETWORK) | frozenset(CREDENTIALED_REMOTE)
_MUST_GUARD = _REMOTE | frozenset(READER_SUPPLIED_PATH)


def test_network_snippets_check_success_before_reading_output():
    """A snippet reports the tool's error instead of raising on ``None``.

    Every snippet that reaches a remote service can fail for a reason outside
    the reader's control — a 5xx, a rate limit, no route to the host, a
    credential the service rejected. Reading ``result.output`` first turns that
    into ``TypeError: 'NoneType' object is not subscriptable`` and discards the
    message the tool produced, so each of these snippets checks
    ``result.success`` and exits with ``result.error``. Needing a credential
    does not exempt a snippet: the host it calls is just as remote.

    A snippet that names a file the reader supplies is checked for the same
    reason. ``/tmp/report.docx`` is not on the reader's machine, so the
    documented call fails there, and the tool's message names the path while
    the ``TypeError`` does not.

    Snippets that build their own input are exempt: they fail only when the
    reader's own machine or arguments are wrong, and the traceback says so.
    """
    problems: list[str] = []
    for heading, code in _gallery_snippets():
        if heading not in _MUST_GUARD:
            continue
        tree = ast.parse(code)
        for name in sorted(_result_names(tree)):
            output_line = _first_output_line(tree, name)
            if output_line is None:
                continue
            guard = _guard_line(tree, name)
            if guard is None:
                problems.append(
                    f"{heading}: reads `{name}.output` (line {output_line}) with no "
                    f"`if not {name}.success:` guard that raises or exits; an upstream "
                    "outage would raise on None and hide the tool's error"
                )
            elif guard > output_line:
                problems.append(
                    f"{heading}: guards `{name}.success` on line {guard}, after "
                    f"`{name}.output` is read on line {output_line}"
                )
    assert not problems, "\n".join(problems)


def test_gallery_snippet_imports_resolve():
    """Every module imported by a snippet is importable (no stale paths)."""
    import importlib

    modules: set[str] = set()
    for _, code in _gallery_snippets():
        for node in ast.walk(ast.parse(code)):
            if isinstance(node, ast.ImportFrom) and node.module:
                if node.module.startswith("effgen"):
                    modules.add(node.module)
    assert modules, "expected effgen imports in the gallery"
    for mod in sorted(modules):
        importlib.import_module(mod)


# ---------------------------------------------------------------------------
# Execution helpers
# ---------------------------------------------------------------------------

# Marker the snippet runner prints when a failed snippet left a ToolResult that
# reports an upstream error, so the caller sees the tool's message even though the
# snippet itself discarded it.
_TOOL_ERROR_MARKER = "SNIPPET_TOOL_ERROR:"

# A snippet that reads ``result.output`` without checking ``result.success`` dies
# on ``None`` and loses the reason. This wrapper runs it unchanged and, if it
# raises, walks the traceback for a ToolResult that failed and prints its error.
_SNIPPET_RUNNER = f'''
import runpy
import sys
import traceback

try:
    runpy.run_path(sys.argv[1], run_name="__main__")
except BaseException as exc:  # noqa: BLE001 - report, then exit non-zero
    traceback.print_exc()
    tb = exc.__traceback__
    while tb is not None:
        try:
            values = list(tb.tb_frame.f_locals.values())
        except Exception:  # noqa: BLE001
            values = []
        for value in values:
            if getattr(value, "success", None) is False and getattr(value, "error", None):
                print("{_TOOL_ERROR_MARKER} " + str(value.error), file=sys.stderr)
        tb = tb.tb_next
    raise SystemExit(1) from None
'''

# Wording used by the adapters and tools for an upstream problem that is not the
# snippet's fault. Same vocabulary as tests/tools/test_semantic_scholar.py, plus
# the refusals a service aims at the network rather than at the request: these
# snippets send no credentials, so such a refusal says where the run happened,
# not that the example is wrong.
_TRANSIENT_UPSTREAM = (
    "rejects unauthenticated api traffic",
    "429",
    "http 500",
    "http 502",
    "http 503",
    "http 504",
    "internal server error",
    "bad gateway",
    "service unavailable",
    "gateway timeout",
    "too many requests",
    "rate-limited",
    "rate limit",
    "timed out",
    "timeout",
    "connection reset",
    "connection error",
    "temporarily unavailable",
)


def _run_snippet(
    code: str,
    tmp_path: Path,
    cwd: Path | None = None,
    attempts: int = 1,
    retry_delay: float = 10.0,
    skip_on_upstream_error: bool = False,
) -> None:
    """Execute a snippet verbatim in a subprocess; /tmp/ paths land in tmp_path.

    *attempts* > 1 re-runs the snippet when it exits non-zero, waiting
    *retry_delay* seconds and doubling that wait each time.

    With *skip_on_upstream_error*, the snippet runs through a wrapper that
    reports the error of any failed ``ToolResult`` still on the traceback, and a
    failure whose reported error is a 429/5xx/timeout from the public service the
    snippet calls becomes a skip. Anything else — a renamed API, a bad argument, a
    changed result shape — still fails, so the check keeps its meaning.
    """
    rewritten = code.replace("/tmp/", f"{tmp_path}/")
    script = tmp_path / "snippet.py"
    script.write_text(rewritten, encoding="utf-8")
    if skip_on_upstream_error:
        runner = tmp_path / "snippet_runner.py"
        runner.write_text(_SNIPPET_RUNNER, encoding="utf-8")
        argv = [sys.executable, str(runner), str(script)]
    else:
        argv = [sys.executable, str(script)]
    for attempt in range(1, attempts + 1):
        proc = subprocess.run(
            argv,
            cwd=cwd or _REPO_ROOT,
            env=os.environ.copy(),
            text=True,
            capture_output=True,
            timeout=110,
            check=False,
        )
        if proc.returncode == 0:
            return
        if attempt < attempts:
            time.sleep(retry_delay * 2 ** (attempt - 1))
    reported = [
        line.split(_TOOL_ERROR_MARKER, 1)[1].strip()
        for line in proc.stderr.splitlines()
        if _TOOL_ERROR_MARKER in line
    ]
    if skip_on_upstream_error:
        for message in reported:
            lowered = message.lower()
            if any(token in lowered for token in _TRANSIENT_UPSTREAM):
                pytest.skip(f"transient upstream error from the snippet's service: {message}")
    raise AssertionError(
        f"snippet failed (rc={proc.returncode}) on {attempts} attempt(s)\n"
        f"--- stdout ---\n{proc.stdout[-2000:]}\n--- stderr ---\n{proc.stderr[-2000:]}"
    )


def test_snippet_runner_reruns_a_flaky_snippet_but_not_a_broken_one(tmp_path):
    """Retries absorb an intermittent failure; a snippet that always fails still fails."""
    marker = tmp_path / "attempts"
    # The marker is resolved next to the generated script, so the /tmp/ rewrite
    # _run_snippet applies to snippet text cannot reach it.
    flaky = (
        "from pathlib import Path\n"
        "p = Path(__file__).parent / 'attempts'\n"
        "n = int(p.read_text()) + 1 if p.exists() else 1\n"
        "p.write_text(str(n))\n"
        "raise SystemExit(0 if n >= 2 else 1)\n"
    )
    _run_snippet(flaky, tmp_path, attempts=3, retry_delay=0)
    assert marker.read_text() == "2"

    with pytest.raises(AssertionError, match="on 3 attempt"):
        _run_snippet("raise SystemExit(1)\n", tmp_path, attempts=3, retry_delay=0)


def test_snippet_runner_separates_an_upstream_outage_from_a_broken_snippet(tmp_path):
    """A 5xx from the snippet's service skips; any other failure still fails."""
    stub = (
        "class ToolResult:\n"
        "    success = False\n"
        "    output = None\n"
        "    error = {error!r}\n"
        "\n"
        "result = ToolResult()\n"
        "print(result.output['results'])\n"
    )
    outage = stub.format(
        error="Tool execution failed: HTTP 500 from https://example.invalid/search: "
              "Internal Server Error"
    )
    with pytest.raises(pytest.skip.Exception, match="transient upstream error"):
        _run_snippet(outage, tmp_path, attempts=1, skip_on_upstream_error=True)

    # The same shape with a non-transient reason is a real failure.
    broken = stub.format(error="Tool execution failed: unknown operation 'serch'")
    with pytest.raises(AssertionError, match="snippet failed"):
        _run_snippet(broken, tmp_path, attempts=1, skip_on_upstream_error=True)

    # And a snippet that breaks on effGen's own API, with no ToolResult involved,
    # is never skipped.
    with pytest.raises(AssertionError, match="snippet failed"):
        _run_snippet(
            "raise AttributeError('SomeTool has no attribute execute_sync')\n",
            tmp_path,
            attempts=1,
            skip_on_upstream_error=True,
        )


def _require_deps(heading: str) -> None:
    for dep in _OPTIONAL_DEPS.get(heading, []):
        pytest.importorskip(dep)


def _make_fixtures(heading: str, tmp_path: Path) -> None:
    """Create the input file a snippet reads, at the path it will see."""
    if heading == "DataFrameTool":
        (tmp_path / "data.csv").write_text("name,score\na,1\nb,2\nc,3\n")
    elif heading == "ImageInfoTool" or heading == "ImageCaptionTool":
        from PIL import Image

        Image.new("RGB", (400, 300), (70, 130, 180)).save(tmp_path / "photo.jpg")
    elif heading == "PDFTool":
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_pdf import PdfPages

        with PdfPages(tmp_path / "paper.pdf") as pdf:
            fig = plt.figure(figsize=(8.5, 11))
            fig.text(0.1, 0.8, "A short sample document.")
            pdf.savefig(fig)
            plt.close(fig)
    elif heading == "DOCXTool":
        import docx

        doc = docx.Document()
        doc.add_heading("Quarterly Report", 0)
        doc.add_paragraph("Revenue grew 12% quarter over quarter.")
        doc.save(str(tmp_path / "report.docx"))
    elif heading == "ExcelTool":
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["name", "score"])
        ws.append(["alpha", 91])
        wb.save(str(tmp_path / "data.xlsx"))
    elif heading == "OCRTool":
        from PIL import Image, ImageDraw

        img = Image.new("RGB", (600, 120), (255, 255, 255))
        ImageDraw.Draw(img).text((20, 40), "Invoice 2026-001", fill=(0, 0, 0))
        img.save(tmp_path / "scan.png")
    elif heading == "AudioTranscribeTool":
        shutil.copy(_AUDIO_FIXTURE, tmp_path / "clip.mp3")


# ---------------------------------------------------------------------------
# Offline snippets
# ---------------------------------------------------------------------------

class TestOfflineSnippets:
    @pytest.mark.cookbook
    @pytest.mark.parametrize("heading", OFFLINE)
    def test_snippet_runs(self, heading, tmp_path):
        _require_deps(heading)
        _make_fixtures(heading, tmp_path)
        blocks = _blocks_for(heading)
        assert blocks, f"no snippet found under heading {heading!r}"
        # Only the first block per heading is a standalone program; later
        # blocks (e.g. the PythonREPL max_sessions line) are fragments.
        _run_snippet(blocks[0], tmp_path)

    @pytest.mark.cookbook
    def test_qr_generate_then_read(self, tmp_path):
        """The QR snippets form a sequence: generate writes, read decodes."""
        _require_deps("QRGenerateTool")
        _require_deps("QRReadTool")
        gen, read = _blocks_for("QRGenerateTool")[0], _blocks_for("QRReadTool")[0]
        _run_snippet(gen + "\n" + read, tmp_path)


# ---------------------------------------------------------------------------
# Network snippets (public services, no credentials)
# ---------------------------------------------------------------------------

class TestNetworkSnippets:
    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.parametrize("heading", sorted(set(NETWORK)))
    def test_snippet_runs(self, heading, tmp_path):
        _require_deps(heading)
        for code in _blocks_for(heading):
            _run_snippet(code, tmp_path, attempts=2, skip_on_upstream_error=True)


# ---------------------------------------------------------------------------
# Gated snippets: credentials / binaries / model keys
# ---------------------------------------------------------------------------

def _docker_available() -> bool:
    if not shutil.which("docker"):
        return False
    probe = subprocess.run(
        ["docker", "ps", "--format", "{{.ID}}"], capture_output=True, timeout=15
    )
    return probe.returncode == 0


class TestGatedSnippets:
    @pytest.mark.cookbook
    @pytest.mark.skipif(not _docker_available(), reason="Docker daemon not reachable")
    def test_docker(self, tmp_path):
        _run_snippet(_blocks_for("DockerTool")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("WOLFRAM_ALPHA_APPID"), reason="WOLFRAM_ALPHA_APPID not set")
    def test_wolfram_alpha(self, tmp_path):
        _run_snippet(_blocks_for("WolframAlphaTool")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("SMTP_HOST"), reason="SMTP not configured")
    def test_email_smtp(self, tmp_path):
        _run_snippet(_blocks_for("EmailSMTPTool")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("IMAP_HOST"), reason="IMAP not configured")
    def test_email_imap(self, tmp_path):
        _run_snippet(_blocks_for("EmailIMAPTool")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("SLACK_WEBHOOK_URL"), reason="SLACK_WEBHOOK_URL not set")
    def test_slack_webhook(self, tmp_path):
        _run_snippet(_blocks_for("SlackWebhookTool")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("DISCORD_WEBHOOK_URL"), reason="DISCORD_WEBHOOK_URL not set")
    def test_discord_webhook(self, tmp_path):
        _run_snippet(_blocks_for("DiscordWebhookTool")[0], tmp_path)

    @pytest.mark.cookbook
    @pytest.mark.skipif(not shutil.which("tesseract"), reason="tesseract not installed")
    def test_ocr(self, tmp_path):
        _make_fixtures("OCRTool", tmp_path)
        _run_snippet(_blocks_for("OCRTool")[0], tmp_path)

    @pytest.mark.cookbook
    def test_audio_transcribe(self, tmp_path):
        # faster-whisper decodes MP3 itself (bundled PyAV); ffmpeg not needed.
        pytest.importorskip("faster_whisper")
        if not _AUDIO_FIXTURE.exists():
            pytest.skip("sample_audio.mp3 fixture not found")
        _make_fixtures("AudioTranscribeTool", tmp_path)
        _run_snippet(_blocks_for("AudioTranscribeTool")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(
        not (os.getenv("OPENAI_API_KEY") or os.getenv("GOOGLE_API_KEY")),
        reason="No vision API key",
    )
    def test_image_caption(self, tmp_path):
        _make_fixtures("ImageCaptionTool", tmp_path)
        _run_snippet(_blocks_for("ImageCaptionTool")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("OPENAI_API_KEY"), reason="OPENAI_API_KEY not set")
    def test_openai_native_agent_constructs(self, tmp_path):
        _run_snippet(_blocks_for("OpenAI Native Tools")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("GOOGLE_API_KEY"), reason="GOOGLE_API_KEY not set")
    def test_gemini_native_agent_constructs(self, tmp_path):
        _run_snippet(_blocks_for("Gemini Native Tools")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("CEREBRAS_API_KEY"), reason="CEREBRAS_API_KEY not set")
    def test_agent_with_tools(self, tmp_path):
        _run_snippet(_blocks_for("Using Tools in an Agent")[0], tmp_path)

    @pytest.mark.live
    @pytest.mark.cookbook
    @pytest.mark.skipif(not os.getenv("CEREBRAS_API_KEY"), reason="CEREBRAS_API_KEY not set")
    def test_presets_construct(self, tmp_path):
        _run_snippet(_blocks_for("Using Presets")[0], tmp_path)
