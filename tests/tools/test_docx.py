"""Tests for DOCXTool - round-trip and error handling."""

from __future__ import annotations

import io

import pytest

from effgen.errors import CorruptDocumentError
from effgen.tools.builtin import DOCXTool

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Create an in-memory DOCX with given paragraphs and optional table."""
    from docx import Document

    doc = Document()
    doc.core_properties.author = "effGen Test"
    doc.core_properties.title = "effGen DOCX Round-trip"
    for para in paragraphs:
        doc.add_paragraph(para)
    if table_rows:
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=0, cols=cols)
        for row_data in table_rows:
            row = table.add_row()
            for i, cell_val in enumerate(row_data):
                row.cells[i].text = cell_val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_bytes():
    return _make_docx(
        ["Hello effGen", "Second paragraph", "DOCX_ROUNDTRIP_TOKEN_99"],
        table_rows=[["Name", "Value"], ["Alpha", "1"], ["Beta", "2"]],
    )


@pytest.fixture
def docx_file(tmp_path, docx_bytes):
    p = tmp_path / "test.docx"
    p.write_bytes(docx_bytes)
    return str(p)


@pytest.fixture
def tool():
    return DOCXTool()


# ---------------------------------------------------------------------------
# Unit tests
# ---------------------------------------------------------------------------

@pytest.mark.unit
def test_docx_tool_metadata():
    t = DOCXTool()
    assert t.metadata.name == "docx"
    assert "docx" in t.metadata.tags


@pytest.mark.unit
def test_corrupt_docx_raises(tool):
    with pytest.raises(CorruptDocumentError):
        import asyncio
        asyncio.run(tool._execute("text", docx_bytes=b"NOT_A_DOCX"))


@pytest.mark.unit
def test_missing_path_raises(tool):
    with pytest.raises(ValueError, match="path"):
        import asyncio
        asyncio.run(tool._execute("text"))


@pytest.mark.unit
def test_unknown_operation_raises(tool, docx_file):
    with pytest.raises(ValueError, match="Unknown operation"):
        import asyncio
        asyncio.run(tool._execute("bogus", path=docx_file))


# ---------------------------------------------------------------------------
# Integration tests
# ---------------------------------------------------------------------------

@pytest.mark.integration
@pytest.mark.asyncio
async def test_docx_text_roundtrip(tool, docx_bytes):
    result = await tool._execute("text", docx_bytes=docx_bytes)
    assert result["success"] is True
    assert "Hello effGen" in result["text"]
    assert "DOCX_ROUNDTRIP_TOKEN_99" in result["text"]


@pytest.mark.integration
@pytest.mark.asyncio
async def test_docx_paragraphs(tool, docx_bytes):
    result = await tool._execute("paragraphs", docx_bytes=docx_bytes)
    assert result["success"] is True
    texts = [p["text"] for p in result["paragraphs"]]
    assert any("Hello effGen" in t for t in texts)
    assert any("DOCX_ROUNDTRIP_TOKEN_99" in t for t in texts)


@pytest.mark.integration
@pytest.mark.asyncio
async def test_docx_tables(tool, docx_bytes):
    result = await tool._execute("tables", docx_bytes=docx_bytes)
    assert result["success"] is True
    assert result["table_count"] == 1
    first_table = result["tables"][0]
    assert first_table["row_count"] == 3
    assert "Name" in first_table["rows"][0]


@pytest.mark.integration
@pytest.mark.asyncio
async def test_docx_metadata(tool, docx_bytes):
    result = await tool._execute("metadata", docx_bytes=docx_bytes)
    assert result["success"] is True
    assert result["metadata"]["author"] == "effGen Test"
    assert result["metadata"]["title"] == "effGen DOCX Round-trip"


@pytest.mark.integration
@pytest.mark.asyncio
async def test_docx_from_file(tool, docx_file):
    result = await tool._execute("text", path=docx_file)
    assert result["success"] is True
    assert "Hello effGen" in result["text"]
